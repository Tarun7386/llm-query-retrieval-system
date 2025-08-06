# services/document_processor.py
from __future__ import annotations

import asyncio
import hashlib
import logging
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from urllib.parse import urlparse

import aiofiles
import aiohttp
import PyPDF2
from docx import Document as DocxDocument
from email import policy
from email.parser import BytesParser
from sqlalchemy.orm import Session

from config import settings
from models.database_models import Document, DocumentChunk
from models.pydantic_models import DocumentType
from utils.helpers import (
    chunk_text_by_sentences,
    clean_text,
    extract_metadata_from_text,
)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Download, extract, chunk, and persist documents."""

    _HTTP_TIMEOUT = aiohttp.ClientTimeout(total=600)  # 10 min for large files

    # ── chunking constants ─────────────────────────────────────────────
    MAX_CHUNK_SIZE = 1_200      # characters
    MIN_CHUNK_SIZE = 200
    OVERLAP_SIZE   = 150

    def __init__(self) -> None:
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
        self._session: Optional[aiohttp.ClientSession] = None

    # ── networking ─────────────────────────────────────────────────────
    async def _client(self) -> aiohttp.ClientSession:
        if self._session is None or self._session.closed:
            self._session = aiohttp.ClientSession(timeout=self._HTTP_TIMEOUT)
        return self._session

    async def download_from_url(self, url: str) -> Tuple[bytes, str]:
        """Return (bytes, filename)."""
        session = await self._client()
        async with session.get(url, max_redirects=5) as resp:
            if resp.status != 200:
                raise ValueError(f"Download failed ({resp.status}) → {url}")

            data = await resp.read()
            fname = self._extract_filename(url, resp.headers)
            logger.info("Downloaded %.2f MB from %s", len(data) / 1_048_576, url)
            return data, fname

    @staticmethod
    def _extract_filename(url: str, headers: Dict) -> str:
        cd = headers.get("content-disposition", "")
        if "filename=" in cd:
            return cd.split("filename=")[1].strip("\"' ")

        path_name = Path(urlparse(url).path).name
        return path_name if "." in path_name else "document.pdf"

    # ── public API ─────────────────────────────────────────────────────
    async def process_document_from_url(self, url: str, db: Session) -> Document:
        data, fname = await self.download_from_url(url)
        mime = mimetypes.guess_type(fname)[0] or "application/octet-stream"
        return await self.process_document(data, fname, mime, db)

    async def process_document(
        self,
        file_bytes: bytes,
        filename: str,
        mime: str,
        db: Session,
    ) -> Document:
        """Persist file, extract text, create chunks, save to DB."""
        doc_type = self._detect_type(filename, mime)

        # save original file
        unique_name = f"{hashlib.md5(file_bytes).hexdigest()}_{filename}"
        file_path = self.upload_dir / unique_name
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(file_bytes)

        document = Document(
            filename=unique_name,
            original_filename=filename,
            file_path=str(file_path),
            file_size=len(file_bytes),
            content_type=mime,
            document_type=doc_type.value,
            processing_status="processing",
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        try:
            raw_text = await self._extract_text(file_path, doc_type)
            text = clean_text(raw_text)

            chunks = self._create_smart_chunks(text, doc_type)
            self._persist_chunks(chunks, document, doc_type, db)

            document.processing_status = "completed"
            document.chunk_count = len(chunks)
            document.document_metadata = extract_metadata_from_text(text)
            db.commit()

            logger.info("Processed %s with %s chunks", filename, len(chunks))
        except Exception as exc:
            document.processing_status = "failed"
            document.processing_error = str(exc)
            db.commit()
            logger.exception("Failed to process %s: %s", filename, exc)
            raise

        return document

    # ── text-extraction helpers ────────────────────────────────────────
    def _detect_type(self, filename: str, mime: str) -> DocumentType:
        ext = Path(filename).suffix.lower()
        if ext == ".pdf" or "pdf" in mime:
            return DocumentType.PDF
        if ext == ".docx" or "wordprocessingml" in mime:
            return DocumentType.DOCX
        if ext == ".txt" or mime.startswith("text/plain"):
            return DocumentType.TXT
        if ext == ".eml" or mime == "message/rfc822":
            return DocumentType.EMAIL
        return DocumentType.PDF  # fallback

    async def _extract_text(self, path: Path, doc_type: DocumentType) -> str:
        if doc_type is DocumentType.PDF:
            return await asyncio.to_thread(self._read_pdf, path)
        if doc_type is DocumentType.DOCX:
            return await asyncio.to_thread(self._read_docx, path)
        if doc_type is DocumentType.TXT:
            async with aiofiles.open(path, "r", encoding="utf-8") as f:
                return await f.read()
        if doc_type is DocumentType.EMAIL:
            return await asyncio.to_thread(self._read_email, path)
        raise ValueError(f"Unsupported type: {doc_type}")

    @staticmethod
    def _read_pdf(path: Path) -> str:
        pages = []
        with path.open("rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for i, page in enumerate(reader.pages):
                pages.append(f"\n--- Page {i+1} ---\n{page.extract_text() or ''}")
        return "\n".join(pages)

    @staticmethod
    def _read_docx(path: Path) -> str:
        doc = DocxDocument(path)
        return "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

    @staticmethod
    def _read_email(path: Path) -> str:
        with path.open("rb") as fh:
            msg = BytesParser(policy=policy.default).parse(fh)
        headers = "\n".join(f"{k}: {msg.get(k, 'N/A')}"
                            for k in ("Subject", "From", "To", "Date"))
        parts: List[str] = []
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    parts.append(part.get_content())
        elif msg.get_content_type() == "text/plain":
            parts.append(msg.get_content())

        return f"{headers}\n---\n" + "\n".join(parts)

    # ── chunking ───────────────────────────────────────────────────────
    def _create_smart_chunks(self, text: str, doc_type: DocumentType) -> List[str]:
        if not text:
            return []
        if doc_type is DocumentType.PDF:
            return self._create_policy_chunks(text)
        return chunk_text_by_sentences(text, self.MAX_CHUNK_SIZE, self.OVERLAP_SIZE)
    # ── revised policy chunker  (drop-in replacement) ─────────────────────
    def _create_policy_chunks(self, text: str) -> List[str]:
        """
        Split insurance PDFs into ~1,200-char chunks with 150-char overlap.
        Handles paragraphs that are themselves longer than MAX_CHUNK_SIZE.
        """
        paragraphs = [p.strip() for p in text.splitlines() if p.strip()]

        chunks: List[str] = []
        buff, current_section = "", ""
        limit, min_size, overlap = self.MAX_CHUNK_SIZE, self.MIN_CHUNK_SIZE, self.OVERLAP_SIZE

        section_markers = {
        "DEFINITIONS", "BENEFITS", "COVERAGE", "EXCLUSIONS", "CONDITIONS",
        "SECTION", "CHAPTER", "ARTICLE", "CLAUSE", "PARAGRAPH",
        "PREMIUM", "CLAIM", "WAITING", "PERIOD", "LIMIT", "DEDUCTIBLE",
        "MATERNITY", "PRE-EXISTING", "HOSPITALISATION", "TREATMENT",
        }

        def flush():
            nonlocal buff
            if len(buff) >= min_size:
                chunks.append(self._finalize_chunk(buff, current_section))
            buff = ""

        for para in paragraphs:
            is_header = any(tok in para.upper() for tok in section_markers)

        # ── process paragraphs longer than the limit ──
            while len(para) > limit:
                head, para = para[:limit], para[limit:]
                if len(buff) + len(head) + 1 > limit:
                    flush()
                buff += ("\n" if buff else "") + head
                flush()

        # ── normal paragraph handling ──
            if len(buff) + len(para) + 1 > limit:
                flush()
                if buff:                      # overlap into new chunk
                    tail = buff[-overlap:]
                    buff = tail + ("\n" if tail else "")
            buff += ("\n" if buff else "") + para

            if is_header and len(para) < 200:
                current_section = para[:100]

        flush()  # final chunk
        logger.info("Created %s policy chunks from %s characters",
                len(chunks), len(text))
        return chunks


    
    @staticmethod
    def _finalize_chunk(chunk_text: str, section_ctx: str) -> str:
        return f"[{section_ctx}]\n{chunk_text}" if section_ctx and section_ctx not in chunk_text else chunk_text

    # ── persist chunks ─────────────────────────────────────────────────
    def _persist_chunks(
        self,
        chunks: List[str],
        document: Document,
        doc_type: DocumentType,
        db: Session,
    ) -> None:
        for i, text in enumerate(chunks):
            page_num = None
            if "--- Page" in text:
                import re
                m = re.search(r"--- Page (\d+) ---", text)
                page_num = m.group(1) if m else None

            chunk = DocumentChunk(
                document_id=document.id,
                chunk_index=i,
                content=text,
                content_hash=hashlib.md5(text.encode()).hexdigest(),
                chunk_metadata={
                    "document_type": doc_type.value,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "chunk_size": len(text),
                    "page_number": page_num,
                    **extract_metadata_from_text(text),
                },
            )
            db.add(chunk)
        db.commit()

    # ── cleanup ────────────────────────────────────────────────────────
    async def delete_document(self, doc_id: str, db: Session) -> bool:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            return False
        Path(doc.file_path).unlink(missing_ok=True)
        db.delete(doc)            # cascades to chunks
        db.commit()
        logger.info("Deleted document %s", doc_id)
        return True
